import os
import pyperclip
import docx
history = ['']
n = 0

def add_text_to_document(doc_path, new_text):
  """Adds a new line with the specified text to a Word document.

  Args:
    doc_path: Path to the Word document.
    new_text: The text to add.
  """

  try:
    doc = docx.Document(doc_path)
    new_paragraph = doc.add_paragraph()
    new_paragraph.add_run(new_text)
    doc.save(doc_path)
    print("Text added successfully.")
  except Exception as e:
    print(f"Error adding text: {e}")

    
def clear_screen():
  """Clears the terminal screen."""
  os.system('cls' if os.name == 'nt' else 'clear')


while True:
   try:
      lastPasted = pyperclip.paste()
      if lastPasted == history[-1]:
         pass
      else:
         history.append(lastPasted)
         if history[0] == '': 
            del history[0]
         clear_screen()
         print(history)
   except KeyboardInterrupt:
      print("interuput called")
      user_input = input("Enter a command (e.g., 'clear'): ")
      if user_input == 'clear' or user_input == 'cls':
         history = ['']
         print("Cleared!")
      else:
         pass


